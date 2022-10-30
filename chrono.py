from docx import Document
from docx.shared import Pt
from docx.shared import *
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert

def createme(name,address, mobno, emailid,linkid, summary, companynames, comyourrole, yearrange, descslist, schoolnames, degreenams, schoolyears, marks, skills,filenamee):
    
    ############################################
    #data to insert
    #name = "Farhan Ahmed"
    #address = "Gwalior, M.P., india"
    #mobno = "9999999999"
    #emailid = "farhan@emailid.com"
    #summary = "Summary goes here, "*5
    #companynames = ["Company1","Company2","Company3"]
    cominfocount = len(companynames)
    #comyourrole = ["Role1","Role2","Role3"]
    #yearrange = ["2010-2012","2012-2014","2014-2016"]
    noofdesc = cominfocount #no of description
    dlist = descslist
    descslist = []
    lenofdlist = len(dlist)
    for l in range(lenofdlist):
        a = dlist[l]
        b = a.split(',')
        descslist.append(b)
        #str(comdesc+'l') = b
        
    #comdesc1 = ["com1desc1","com1desc2","com1desc3"] 
    #comdesc2 = ["com2desc1","com2desc2","com2desc3"]
    #comdesc3 = ["com3desc1","com3desc2","com3desc3"]
    #ncomdesc1 = len(comdesc1)
    #ncomdesc2 = len(comdesc2)
    #ncomdesc3 = len(comdesc3)
    #descslist = [comdesc1,comdesc2,comdesc3]
    dot = u'\u2022'
    #schoolnames = ["School1","School2","School3"]
    schoolcount = len(schoolnames)
    #degreenams = ["Degree1","Degree2","Degree3"]
    #schoolyears = ["2000-2004","2004-2008","2008-2010"]
    #marks = ["80%","90%","95%"]
    #skills = ["skill1","skill2","skill3"]
    skillscount = len(skills)
    ############################################
    ############################################
    #size color style of font
    #colors
    namecolor = RGBColor(0,102,255)
    textcolor= RGBColor(0,0,102) #regular texts
    cnameroleyearcolor = RGBColor(51,102,255) #company name, role, year
    #sizes
    namesize = Pt(30)
    textsize = Pt(11) #regular texts
    addmobesize = Pt(12) #address, mobile no, email id 
    cnameroleyearsize = Pt(14) #company name, role, year
    schoolnamesize = Pt(14) #school name
    #style family
    textfamily = "Cambria"


    ############################################
    # source Template
    document = Document('C:\\Users\\lenevo\\Desktop\\Resume with Flask\\frontend\\chrnologicalR\\chronoTemp.docx')
    #C:\Users\lenevo\Desktop\Resume with Flask\frontend\chrnologicalR
    ############################################
    np = len(document.paragraphs)
    i = 0
    for obj in document.paragraphs:
        if obj.text == '':
            r = len(document.tables[i].rows) #Find the no. of rows in the table
            for row in range(r):
                c = len(document.tables[i].rows[row].cells) #Find the no. of cells in the table
                txt = document.tables[i].rows[row].cells[0].text
                #start of matching and inserting
                if txt == "Your Name":
                    #erase the text
                    document.tables[i].rows[row].cells[0].paragraphs[0].clear()
                    #add info with font size and style
                    par = document.tables[i].rows[row].cells[0].add_paragraph()
                    par.alignment = WD_ALIGN_PARAGRAPH.CENTER 
                    para = par.add_run(name+"\n")
                    para.font.size = namesize
                    para.font.bold = True
                    para.font.color.rgb = namecolor
                    para.font.name = textfamily
                    para1 = par.add_run(address+"\n")
                    para1.font.size = addmobesize
                    para1.font.color.rgb = textcolor
                    para2 = par.add_run(mobno+"\n")
                    para2.font.size = addmobesize
                    para2.font.color.rgb = textcolor
                    para3 = par.add_run(emailid+"\n")
                    para3.font.size = addmobesize
                    para3.font.color.rgb = textcolor
                    para9 = par.add_run(linkid+"\n")
                    para9.font.size = addmobesize
                    para9.font.color.rgb = textcolor
                    
                elif txt == "Summary":
                    para4 = document.tables[i].rows[row].cells[0].paragraphs[0].clear()                
                    para41 =para4.add_run(summary)
                    para41.font.size = textsize
                    para41.font.name = textfamily
                    para41.font.color.rgb = textcolor
                elif txt == "company":
                    document.tables[i].rows[row].cells[0].paragraphs[0].clear()
                    for company in range(cominfocount):
                        para5 = document.tables[i].rows[row].cells[0].add_paragraph()
                        para51 = para5.add_run(companynames[company]+"\n")
                        para52 = para5.add_run(comyourrole[company]+"\n")
                        para53 = para5.add_run(yearrange[company])
                        para51.font.size, para51.font.name, para51.font.color.rgb = cnameroleyearsize, textfamily, cnameroleyearcolor
                        para52.font.size, para52.font.name, para52.font.color.rgb = cnameroleyearsize, textfamily, cnameroleyearcolor
                        para52.font.size, para52.font.name, para53.font.color.rgb = cnameroleyearsize, textfamily, cnameroleyearcolor
                        para6 = document.tables[i].rows[row].cells[0].add_paragraph()
                        noofdesc = len(descslist[company])
                        for desc in range(noofdesc):
                            if desc != noofdesc-1:
                                para61 = para6.add_run(dot+" "+descslist[company][desc]+"\n")
                                para61.font.size = textsize
                                para61.font.name = textfamily
                                para61.font.color.rgb = textcolor
                                #add right indentation
                                para6.paragraph_format.left_indent = Inches(0.2) 
                            else:
                                para61 = para6.add_run(dot+" "+descslist[company][desc])
                                para61.font.size = textsize
                                para61.font.name = textfamily
                                para61.font.color.rgb = textcolor
                                #add right indentation
                                para6.paragraph_format.left_indent = Inches(0.2)
                        
                        document.tables[i].rows[row].cells[0].add_paragraph()
                elif txt == "eduinfo":
                    document.tables[i].rows[row].cells[0].paragraphs[0].clear()
                    for school in range(schoolcount):
                        para7 = document.tables[i].rows[row].cells[0].add_paragraph()
                        para71 = para7.add_run(schoolnames[school]+"\n")
                        para72 = para7.add_run(degreenams[school]+"\n")
                        para73 = para7.add_run(schoolyears[school]+"\n")
                        para74 = para7.add_run(marks[school]+"\n")
                        para71.font.size, para71.font.name, para71.font.color.rgb, para71.bold, para71.italic = schoolnamesize, textfamily, textcolor ,True, True
                        para72.font.size, para72.font.name, para72.font.color.rgb, para72.bold, para72.italic = schoolnamesize, textfamily, textcolor ,True, True
                        para73.font.size, para73.font.name, para73.font.color.rgb, para73.bold, para73.italic = schoolnamesize, textfamily, textcolor ,True, True
                        para74.font.size, para74.font.name, para74.font.color.rgb, para74.bold, para74.italic = schoolnamesize, textfamily, textcolor ,True, True
                elif txt == "skills":
                    #erase the text
                    para8 = document.tables[i].rows[row].cells[0].paragraphs[0].clear()
                    for skill in range(skillscount):
                        para81 = para8.add_run(dot+ " "+skills[skill]+"\n")
                        para81.font.size = textsize
                        para81.font.name = textfamily
                        para81.font.color.rgb = textcolor
    
                    
    #filenamee = input("Enter the name of the file: ")
    document.save('C:\\Users\\lenevo\\Desktop\\Resume with Flask\\frontend\\static\\files\\'+filenamee+'.docx')
    convert('C:\\Users\\lenevo\\Desktop\\Resume with Flask\\frontend\\static\\files\\'+filenamee+'.docx', 'C:\\Users\\lenevo\\Desktop\\Resume with Flask\\frontend\\static\\pdf\\'+filenamee+'.pdf')
    #C:\Users\lenevo\Desktop\Resume with Flask\frontend\result24